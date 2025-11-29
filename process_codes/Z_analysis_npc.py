# -*- coding: utf-8 -*-
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
file_path = BASE_DIR / "data" / "all-top5.csv"
output_path = BASE_DIR / "data" / "Supplementary_Table1.docx"

df = pd.read_csv(file_path)

field_map = {
    1: "Chemistry",
    2: "Pharmacology",
    3: "Physics",
    4: "BMA"
}

def get_summary(data, name):
    q = data["new_phrase_comb"].quantile([0.5, 0.75, 0.9, 0.95, 0.99]).to_dict()
    return {
        "Field": name,
        "Count": len(data),
        "Mean": round(data["new_phrase_comb"].mean(), 2),
        "Median": int(q[0.5]),
        "75th": int(q[0.75]),
        "90th": int(q[0.9]),
        "95th": int(q[0.95]),
        "99th": int(q[0.99]),
        "Max": int(data["new_phrase_comb"].max())
    }

results = []
results.append(get_summary(df, "Overall"))

for f in sorted(df["field"].unique()):
    sub = df[df["field"] == f]
    results.append(get_summary(sub, field_map.get(f, f)))

summary_df = pd.DataFrame(results)

doc = Document()
doc.add_heading("Supplementary Table 1. Distribution of new phrase combinations by field", level=1)


table = doc.add_table(rows=1, cols=len(summary_df.columns))
table.style = "Table Grid"

hdr_cells = table.rows[0].cells
for i, col in enumerate(summary_df.columns):
    hdr_cells[i].text = col

for _, row in summary_df.iterrows():
    row_cells = table.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = str(val)

doc.save(output_path)
