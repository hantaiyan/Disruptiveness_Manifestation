# -*- coding: utf-8 -*-
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches

# ========= 参数配置 =========
BASE = Path(__file__).resolve().parent.parent

BASE_DIR = BASE / "results" 
years = [1999, 2004, 2009]
fields = ["Chemistry", "Pharmacology", "Physics", "BMA"]

# ========= 输出路径 =========
word_output = BASE / "data" / "Supplementary_Table7.docx"

# ========= 主程序 =========
all_results = []

for year in years:
    for field in fields:
        file_path = BASE_DIR /  "top5" / f"{year}{field}_focal_DI_y1_to_y15.csv"
        if not file_path.exists():
            print(f"⚠️ 未找到文件: {file_path}")
            continue

        df = pd.read_csv(file_path)

        # 提取 DI_y3 ~ DI_y15 的列
        di_cols = [col for col in df.columns if col.startswith("DI_y") and col[4:].isdigit()]
        di_cols = [col for col in di_cols if 3 <= int(col[4:]) <= 15]
        if not di_cols:
            print(f"❌ {file_path} 不包含 DI_y3–DI_y15 列")
            continue

        # 计算每列的 top5% 阈值
        thresholds = {"year": year, "field": field}
        for col in di_cols:
            val = df[col].quantile(0.95)
            thresholds[f"{col}_threshold"] = round(val, 3)  # 保留三位小数

        all_results.append(thresholds)

# ========= 汇总输出 =========
result_df = pd.DataFrame(all_results)

# 按列顺序整理（year, field, DI_y3_threshold ... DI_y15_threshold）
cols_order = ["year", "field"] + [f"DI_y{i}_threshold" for i in range(3, 16)]
result_df = result_df.reindex(columns=cols_order)

# ========= 写入 Word =========
doc = Document()
doc.add_heading("Top 5% Thresholds (DI_y3–DI_y15)", level=1)

# 添加表格
rows, cols = result_df.shape
table = doc.add_table(rows=rows + 1, cols=cols)
table.style = "Table Grid"

# 写入表头
for j, col_name in enumerate(result_df.columns):
    table.cell(0, j).text = col_name

# 写入数据
for i in range(rows):
    for j, col_name in enumerate(result_df.columns):
        val = result_df.iloc[i, j]
        if isinstance(val, float):
            val = f"{val:.3f}"  # 保留三位小数
        table.cell(i + 1, j).text = str(val)

# 自动调整列宽（可选）
for col in table.columns:
    for cell in col.cells:
        cell.width = Inches(1)

doc.save(word_output)
print(f"✅ 已保存为 Word 文件：{word_output}")
