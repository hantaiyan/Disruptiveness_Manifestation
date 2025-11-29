# -*- coding: utf-8 -*-
from pathlib import Path
import pandas as pd
from docx import Document

BASE = Path(__file__).resolve().parent.parent

file_path = BASE_DIR = OUT_PATH = BASE / "data/all-top5.csv"
df = pd.read_csv(file_path)


field_map = {
    1: "Chemistry",
    2: "Pharmacology",
    3: "Physics",
    4: "BMA"
}

# ---- doc init ----
doc = Document()

# ordering O1: field顺序固定，然后每field内部 year ascending
for fid in [1,2,3,4]:
    for yr in sorted(df["year"].unique()):
        sub = df[(df.field==fid) & (df.year==yr)]
        if sub.empty:
            continue
        
        field_name = field_map[fid]
        doc.add_heading(f"{field_name}-{yr}", level=2)
        
        total = sub.shape[0]
        vc = sub["delay_year"].value_counts().sort_index()
        
        table = doc.add_table(rows=1, cols=5)
        hdr = ["Field-Year","Time window","Delay year","Sample Size","Cumulative proportion"]
        for i,h in enumerate(hdr):
            table.rows[0].cells[i].text = h
        
        cum = 0
        for dy, cnt in vc.items():
            cum += cnt
            time_window = dy + 3
            pct = f"{(cum/total*100):.3f}%"
            
            row = table.add_row().cells
            row[0].text = f"{field_name}-{yr}"
            row[1].text = str(time_window)
            row[2].text = str(dy)
            row[3].text = str(cnt)
            row[4].text = pct
        
        # spacing S2：每个 field-year table 之后 page break
        doc.add_page_break()

out_path = BASE / "data/Supplementary_Table6.docx"
doc.save(out_path)
print("✅ Done: ", out_path)
